import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_full_details_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Firma-KI: Comprehensive Technical & Project Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "Firma-KI is an advanced, enterprise-grade AI proxy gateway, security firewall, and semantic network optimizer. "
        "It is uniquely designed to sit seamlessly between internal corporate infrastructure and commercial Large Language Models "
        "(LLMs) like OpenAI, Anthropic, or DeepSeek. Its core mission is to enable safe, compliant, and extraordinarily cost-effective AI adoption."
    )
    
    doc.add_heading('2. The Pure GDPR Shield (Zero-Trust PII Anonymizer)', level=1)
    doc.add_paragraph(
        "To comply with strict data privacy laws (such as GDPR/DSGVO), Firma-KI employs a deterministic, zero-trust middleware layer."
    )
    doc.add_paragraph("Architecture - The 3-Step Vault:", style='Heading 2')
    doc.add_paragraph(
        "1. Intercept & Mask (Outbound): Local interception of queries with fast NLP and regex patterns to mask PII (Names, Addresses, "
        "IBANs, Project Codes) via cryptographic placeholders (e.g., [PER_01]).", style='List Number'
    )
    doc.add_paragraph(
        "2. Sanitized Processing: The commercial cloud AI processes prompts strictly via placeholders, remaining entirely blind to sensitive data.", style='List Number'
    )
    doc.add_paragraph(
        "3. Re-Injection (Inbound): The gateway maps returned placeholders back to their real-world values before presenting the response to the user.", style='List Number'
    )
    
    doc.add_heading('3. The 3-Stage NEX Bytecode Pipeline', level=1)
    doc.add_paragraph(
        "Firma-KI dramatically reduces API costs (up to 97%) using a multi-layered compression format known as NEX Bytecode."
    )
    doc.add_paragraph(
        "Stage 1 (Compressor & PII Filter): A cheap/local AI model strips PII and compresses bloated prompts into ultra-dense NEX Bytecode.", style='List Number'
    )
    doc.add_paragraph(
        "Stage 2 (The Middle AI): The main commercial LLM processes the ultra-dense bytecode, minimizing expensive token ingestion.", style='List Number'
    )
    doc.add_paragraph(
        "Stage 3 (Human Translator): An expander model translates returning bytecode back into a naturally formulated human response while preserving code and JSON structures.", style='List Number'
    )

    doc.add_heading('4. Key Features & Modules', level=1)
    doc.add_paragraph("• Financial Efficiency Tracking: Live telemetry mapping token payloads and visual savings calculations.", style='List Bullet')
    doc.add_paragraph("• Total Transparency Audit: A developer dashboard exposing Stage 1, Stage 2, and Stage 3 payloads.", style='List Bullet')
    doc.add_paragraph("• Two-Way GDPR Masking: Advanced context-aware filtering instead of simple static regex.", style='List Bullet')
    doc.add_paragraph("• Role-Based Access Control: Granular permissions for organizational members.", style='List Bullet')
    doc.add_paragraph("• File Analysis Hub (RAG): Secure upload and parsing of PDFs for Retrieval-Augmented Generation.", style='List Bullet')
    doc.add_paragraph("• Secure API Streaming: Drop-in replacement for OpenAI endpoints via proxying.", style='List Bullet')

    doc.add_heading('5. System Architecture & Tech Stack', level=1)
    doc.add_paragraph(
        "The Firma-KI platform is built as a highly concurrent system:"
    )
    doc.add_paragraph("Backend Framework: Asynchronous Django 5.x channels serving as the backbone REST API.", style='List Bullet')
    doc.add_paragraph("Database state: Managed securely through PostgreSQL/SQLite.", style='List Bullet')
    doc.add_paragraph("Frontend UI: Rendered fast and responsibly with TailwindCSS and Alpine.js.", style='List Bullet')

    doc.save('FirmaKI_Full_Documentation.docx')
    print("Created FirmaKI_Full_Documentation.docx")


def create_investor_one_pager():
    doc = Document()
    
    # Header
    title = doc.add_heading('Firma-KI: Enterprise AI Gateway', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Secure. Private. Cost-Effective AI Operations.')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    # Core Pitch
    doc.add_heading('The Problem', level=2)
    doc.add_paragraph(
        "Enterprises cannot fully adopt breakthrough AI models (OpenAI, Anthropic) because sending proprietary corporate data "
        "and PII to the cloud violates strict GDPR regulations. Simultaneously, bloated API costs (up to $75 per 1M tokens) limit "
        "scalable production operations."
    )
    
    doc.add_heading('The Solution', level=2)
    doc.add_paragraph(
        "Firma-KI is an on-premise Semantic Firewall and API Gateway that intercepts all outbound AI traffic. "
        "It acts as a drop-in proxy (compatible with legacy code) that completely anonymizes sensitive data and compresses prompts "
        "by up to 90%, all before information ever hits a public cloud."
    )
    
    # Key Pillars
    doc.add_heading('Core Value Proposition', level=2)
    
    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run("Absolute GDPR Compliance (Zero-Trust): ").bold = True
    p1.add_run("Uses local NLP to detect and replace PII (Names, IBANs, internal IDs) with cryptographic placeholders before exporting data to external AIs. Upon returning, data is flawlessly injected back.")

    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run("97% API Cost Reduction (NEX Bytecode): ").bold = True
    p2.add_run("Firma-KI semantically crushes verbose prompts into dense bytecode instructions (the NEX Pipeline), dramatically slashing the token burn rate at commercial providers.")

    p3 = doc.add_paragraph(style='List Bullet')
    p3.add_run("Zero-Friction Deployment: ").bold = True
    p3.add_run("Developers simply change a single line of code (Base URL) to point to Firma-KI. No app rewrites are required. Full transparent audit dashboard provided natively.")
    
    # Market & ROI
    doc.add_heading('Return on Investment', level=2)
    doc.add_paragraph(
        "Enterprises achieve full regulatory compliance instantly. Moreover, the token savings generated by the NEX Pipeline "
        "often absorb the cost of the Firma-KI license within months. It enables the immense intelligence of cloud AI while retaining "
        "the security posture of a private airgap."
    )
    
    doc.save('FirmaKI_Investor_One_Pager.docx')
    print("Created FirmaKI_Investor_One_Pager.docx")

if __name__ == '__main__':
    create_full_details_doc()
    create_investor_one_pager()
