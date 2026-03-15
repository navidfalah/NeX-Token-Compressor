import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_investor_pitch():
    doc = docx.Document()
    
    # Define styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Helvetica'
    font.size = Pt(11)
    
    # Title Page
    doc.add_heading('Firma-KI Nex', 0)
    subtitle = doc.add_paragraph('Enterprise AI API Gateway & Compression Engine')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(79, 70, 229) # Indigo
    
    doc.add_paragraph('\n\n\n')
    pitch = doc.add_paragraph('Investor Presentation & Technical Overview')
    pitch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pitch.runs[0].font.size = Pt(14)
    pitch.runs[0].font.italic = True
    
    doc.add_page_break()
    
    # Problem Statement
    doc.add_heading('1. The Problem: Scaling Generative AI Cost & Latency', level=1)
    doc.add_paragraph(
        "As enterprises rapidly integrate Generative AI into their workflows, they face two critical scaling bottlenecks: "
        "exponentially accelerating API token costs and massive latency delays. "
        "LLMs (Large Language Models) require immense context windows for reasoning, meaning businesses pay per character "
        "for white spaces, formatting, redundancies, and boilerplate that the AI natively doesn't even need to process."
    )
    
    # Solution
    doc.add_heading('2. The Solution: Firma-KI Nex', level=1)
    doc.add_paragraph(
        "Firma-KI Nex is a proprietary middleware proxy gateway that interfaces seamlessly between an enterprise application and "
        "any LLM (DeepSeek, Gemini, OpenAI, Anthropic). By aggressively and deterministically compressing prompts before they "
        "hit paid external APIs, Firma-KI cuts token usage by 60–90% while returning identical AI outputs."
    )
    
    doc.add_heading('3. Core Technical Value Drivers', level=2)
    
    doc.add_heading('A. The NEX Pipeline & Algorithmic Compression', level=3)
    doc.add_paragraph(
        "Firma-KI intercepts data passing to the LLM and processes it through specialized algorithms (both code and natural language). "
        "It eliminates AST syntax, removes type annotations, deduplicates text mathematically, and condenses human-readable text into "
        "dense, machine-native mathematical representations. It is entirely deterministic; there is no stochastic data loss."
    )
    
    doc.add_heading('B. Dynamic Cascade Intelligence Routing', level=3)
    doc.add_paragraph(
        "Not all requests require the same level of intelligence. Firma-KI evaluates prompt complexity natively. Simple extractions "
        "or summaries are routed immediately to low-cost, high-speed tier 1 models (like DeepSeek or Llama). Highly complex logic generation "
        "is routed to Tier 2 powerhouses (like Gemini 2.5 Flash or GPT-4o). Operations are fail-safe and load-balanced automatically."
    )
    
    doc.add_heading('C. Absolute Auditability & Security', level=3)
    doc.add_paragraph(
        "Every byte of data is accounted for. Firma-KI tracks original token counts, compressed token counts, the dynamically chosen "
        "AI provider, and the exact mathematical USD cost savings for every single request. Using the beautiful Glassmorphic dashboard, "
        "administrators can monitor fleet-wide metrics, export PDF reports, and enforce rate limits instantly."
    )
    
    doc.add_page_break()
    
    # Architecture
    doc.add_heading('4. Architecture Highlights', level=1)
    doc.add_paragraph(
        "Developed primarily on the asynchronous FastAPI Python framework, Firma-KI runs a blazing-fast proxy. The infrastructure integrates:"
    )
    doc.add_paragraph("• A resilient Router Engine tracking concurrent connections.", style='List Bullet')
    doc.add_paragraph("• Seamless RBAC (Role-Based Access Control) bridging administrators with individual developers.", style='List Bullet')
    doc.add_paragraph("• A multi-tenant SQLite/PostgreSQL architecture parsing telemetry locally for absolute GDPR compliance.", style='List Bullet')
    doc.add_paragraph("• Fully decoupled input and output pipelines for absolute stability against hallucinated AI responses.", style='List Bullet')
    
    # Financial Impact
    doc.add_heading('5. Investment & Financial Impact', level=1)
    doc.add_paragraph(
        "Firma-KI replaces raw API consumption with efficient, optimized scaling. An organization spending $100k annually on LLM tokens "
        "stands to instantly cut operating costs down to $25k-$40k. Firma-KI shifts Generative AI from an experimental operational cost "
        "into a highly scalable infrastructure layer."
    )
    
    # Finish
    doc.add_paragraph('\n\n\n\nConfidential & Proprietary - Firma-KI Enterprise')
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].runs[0].font.size = Pt(9)
    doc.paragraphs[-1].runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.save('FirmaKI_Nex_Investor_Presentation.docx')

if __name__ == "__main__":
    create_investor_pitch()
